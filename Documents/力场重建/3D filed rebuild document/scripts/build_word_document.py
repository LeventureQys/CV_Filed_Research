from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "3D曲面力场重建方法.md"
OUTPUT = ROOT / "3D曲面力场重建方法（模量科技 甄聪）.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
TEXT_GRAY = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_run_font(run, east_asia="微软雅黑", latin="Arial", size=None, bold=None, color=None) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键选择“更新域”以刷新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    paragraph._p.append(begin)
    paragraph._p.append(code)
    paragraph._p.append(separate)
    paragraph._p.append(text)
    paragraph._p.append(end)


def add_page_number(paragraph) -> None:
    add_field(paragraph, " PAGE ")


def add_inline_markdown(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\$[^$]+?\$)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="等线", latin="Consolas", size=9.5, color=RGBColor(31, 78, 121))
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EAF2F8")
            run._element.get_or_add_rPr().append(shading)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="Cambria Math", latin="Cambria Math", size=10.5)
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def setup_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after in (
        ("Title", 26, 0, 18),
        ("Heading 1", 17, 18, 8),
        ("Heading 2", 14, 15, 6),
        ("Heading 3", 12, 12, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121) if name != "Title" else RGBColor(25, 55, 85)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    caption.font.size = Pt(9)
    caption.font.color.rgb = TEXT_GRAY
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation.font.size = Pt(11)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(7)
    equation.paragraph_format.keep_together = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.6)
    code.paragraph_format.right_indent = Cm(0.6)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0


def setup_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)


def add_cover(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("技术方法文档")
    set_run_font(run, size=12, bold=True, color=RGBColor(31, 78, 121))
    eyebrow.paragraph_format.space_after = Pt(20)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.25
    run = title.add_run("Wendland 紧支撑核的\n3D 曲面力场重建方法")
    set_run_font(run, size=26, bold=True, color=RGBColor(25, 55, 85))

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(20)
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, size=12, color=RGBColor(91, 155, 213))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(12)
    run = subtitle.add_run("离散曲面采样 · 测地距离 · 归一化局部插值")
    set_run_font(run, size=12, color=TEXT_GRAY)

    for _ in range(7):
        document.add_paragraph()
    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(8)
    run = author.add_run("模量科技")
    set_run_font(run, size=14, bold=True, color=RGBColor(31, 78, 121))
    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("甄聪")
    set_run_font(run, size=13, bold=True)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(18)
    run = date.add_run("2026 年 7 月")
    set_run_font(run, size=10.5, color=TEXT_GRAY)

    document.add_page_break()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(18)
    run = heading.add_run("目  录")
    set_run_font(run, size=18, bold=True, color=RGBColor(31, 78, 121))
    headings = []
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^##\s+(.+)$", line.strip())
        if match:
            headings.append(match.group(1))
    for number, title in enumerate(headings, start=1):
        toc = document.add_paragraph()
        toc.paragraph_format.left_indent = Cm(0.8)
        toc.paragraph_format.right_indent = Cm(0.8)
        toc.paragraph_format.space_after = Pt(5)
        number_run = toc.add_run(f"{number:02d}  ")
        set_run_font(number_run, size=10, bold=True, color=RGBColor(31, 78, 121))
        title_run = toc.add_run(title)
        set_run_font(title_run, size=10.5)
    document.add_page_break()


def configure_headers_footers(document: Document) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("3D 曲面力场重建方法  |  模量科技 甄聪")
    set_run_font(run, size=8.5, color=TEXT_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—  ")
    set_run_font(run, size=8.5, color=TEXT_GRAY)
    add_page_number(footer)
    run = footer.add_run("  —")
    set_run_font(run, size=8.5, color=TEXT_GRAY)


def clean_equation(lines: list[str]) -> str:
    text = " ".join(line.strip() for line in lines)
    replacements = {
        r"\mathbf ": "",
        r"\mathbb R": "ℝ",
        r"\qquad": "    ",
        r"\le": "≤",
        r"\ge": "≥",
        r"\in": "∈",
        r"\sum": "∑",
        r"\sqrt": "√",
        r"\frac": "",
        r"\operatorname": "",
        r"\left": "",
        r"\right": "",
        r"\min": "min",
        r"\pi": "π",
        r"\phi": "φ",
        r"\theta": "θ",
        r"\lambda": "λ",
        r"\alpha": "α",
        r"\cdot": "·",
        r"\times": "×",
        r"\text": "",
        r"\\": "    ",
        r"\begin{cases}": "{ ",
        r"\end{cases}": "",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"([A-Za-z])_\{([^}]+)\}", r"\1_\2", text)
    text = re.sub(r"([A-Za-z])_([A-Za-z0-9]+)", r"\1_\2", text)
    text = re.sub(r"\^\{([^}]+)\}", r"^\1", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("&", "   ").replace("~", " ")
    return re.sub(r"\s+", " ", text).strip()


def add_picture(document: Document, path: Path, caption: str) -> None:
    from PIL import Image

    with Image.open(path) as image:
        width, height = image.size
    max_width = Inches(6.25)
    max_height = Inches(5.8)
    ratio = width / height
    picture_width = max_width
    if max_width / ratio > max_height:
        picture_width = max_height * ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=picture_width)
    caption_paragraph = document.add_paragraph(style="Figure Caption")
    add_inline_markdown(caption_paragraph, caption)


def add_markdown_table(document: Document, rows: list[str]) -> None:
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in parsed[1]):
        parsed.pop(1)
    table = document.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(parsed):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(
                run,
                size=9,
                bold=row_index == 0,
                color=RGBColor(255, 255, 255) if row_index == 0 else None,
            )
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        paragraph = cell.add_paragraph(style="Code Block")
        run = paragraph.add_run(line or " ")
        set_run_font(run, east_asia="等线", latin="Consolas", size=9)
    document.add_paragraph()


def build_document() -> None:
    document = Document()
    setup_page(document)
    setup_styles(document)
    configure_headers_footers(document)
    add_cover(document)

    document.core_properties.title = "Wendland 紧支撑核的 3D 曲面力场重建方法"
    document.core_properties.author = "模量科技 甄聪"
    document.core_properties.subject = "三维曲面散点场重建技术方法"
    document.core_properties.keywords = "Wendland, 曲面力场, 场重建, 测地距离, 模量科技"

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    first_title = True
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            if level == 1 and first_title:
                first_title = False
            else:
                paragraph = document.add_paragraph(style=f"Heading {min(level - 1, 3)}")
                add_inline_markdown(paragraph, heading.group(2))
            index += 1
            continue

        image = re.match(r"^!\[(.+?)\]\((.+?)\)$", stripped)
        if image:
            add_picture(document, ROOT / image.group(2), image.group(1))
            index += 1
            continue

        if stripped == "$$":
            equation_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation_lines.append(lines[index])
                index += 1
            paragraph = document.add_paragraph(style="Equation")
            run = paragraph.add_run(clean_equation(equation_lines))
            set_run_font(run, east_asia="Cambria Math", latin="Cambria Math", size=11)
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        list_match = re.match(r"^([-*]|\d+\.)\s+(.+)$", stripped)
        if list_match:
            marker, content = list_match.groups()
            style = "List Bullet" if marker in ("-", "*") else "List Number"
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline_markdown(paragraph, content)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if re.match(r"^(#{1,6})\s+", candidate) or candidate == "$$" or candidate.startswith("```"):
                break
            if candidate.startswith("|") or candidate.startswith("!["):
                break
            if re.match(r"^([-*]|\d+\.)\s+", candidate):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        add_inline_markdown(paragraph, "".join(paragraph_lines))

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
